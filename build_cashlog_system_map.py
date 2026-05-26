from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "cashlog-current-system-map.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D8E2EE", size=8):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_indent(table, dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color_hex, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color_hex)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("ClinicSnap Cash Log\nCurrent System Map")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    meta_run = meta.add_run(
        "Purpose: quick technical reference for how the cash entry tool works today\n"
        "Last documented: May 26, 2026"
    )
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor(0x55, 0x65, 0x77)


def add_stack_table(document):
    document.add_paragraph("Technology Stack", style="Heading 2")
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table, color="D8E2EE", size=8)
    set_table_indent(table, 120)

    widths = [1.2, 1.35, 3.95]
    headers = ["Layer", "Language / Tech", "Current Use"]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(headers[idx])
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    rows = [
        ("Frontend", "HTML / CSS / JS", "Single-file browser app with custom UI, receipt printing, shared dashboard, history, and shift close report."),
        ("Backend", "Google Apps Script", "Receives form submissions, reads Google Sheets, and serves shared dashboard/history/report data."),
        ("Database", "Google Sheets", "Stores all transaction rows in the FormData tab and drives shared reporting."),
        ("Assets", "PNG", "StethoMD and Help360 branding images used by the front-end interface."),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = Pt(9.5)


def add_flow_and_features(document):
    document.add_paragraph("Current Flow", style="Heading 2")
    steps = [
        "Staff enters department, staff name, patient ID, optional receipt patient name, service, amount, payment method, and note.",
        "The browser auto-generates payment date and payment time, then sends the entry to the Apps Script web app by POST.",
        "Apps Script appends the row into the FormData sheet and shared read endpoints pull totals, history, and shift-close data back from that sheet.",
        "The browser builds the receipt locally, supports reprint from shared history, and can export or print shift-close summaries."
    ]
    for text in steps:
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"• {text}")
        run.font.size = Pt(10)

    document.add_paragraph("Key Functional Notes", style="Heading 2")
    notes = [
        "Services and prices are still hardcoded in the HTML; fixed-price services lock the amount while manual entries stay editable.",
        "Shared read features use JSONP script-tag requests to Apps Script, while form submission uses fetch POST with no-cors.",
        "Receipt history and shift close reports are shared across staff because they are read from the common Google Sheet.",
        "Current Google Sheet columns are: Timestamp, Staff Name, Department, Patient ID, Service, Amount Collected, Payment Date, Payment Method, Note, Payment Time, Patient Name."
    ]
    for text in notes:
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"• {text}")
        run.font.size = Pt(10)


def add_endpoints_table(document):
    document.add_paragraph("Active Endpoints and Reports", style="Heading 2")
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table, color="D8E2EE", size=8)
    set_table_indent(table, 120)

    widths = [1.85, 4.65]
    headers = ["Apps Script Action", "Current Result"]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(headers[idx])
        run.font.bold = True
        run.font.size = Pt(10)

    rows = [
        ("POST /exec", "Creates a new payment entry row in Google Sheets."),
        ("action=dashboard", "Returns today's shared cash total, Zelle total, and entry count."),
        ("action=history", "Returns shared receipt history entries used for view and reprint."),
        ("action=shiftReport", "Returns grouped staff totals, department totals, and detailed transactions for the selected date."),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = Pt(9.5)


def build_document():
    document = Document()
    style_document(document)
    add_title(document)
    add_stack_table(document)
    add_flow_and_features(document)
    add_endpoints_table(document)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
